#!/usr/bin/env python3
"""Genera el reporte consolidado en español con doce páginas deliberadas."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


HERE = Path(__file__).resolve().parent
IMAGES = HERE / "images"
OUTPUT = HERE / "reporte_consolidado_experimentos_es.docx"
BLUE = "24577A"
ORANGE = "C87525"
INK = "243746"
MUTED = "61717F"
PALE = "EAF1F6"
WHITE = "FFFFFF"


def font(run, size=9.2, bold=False, color=INK, italic=False):
    run.font.name = "Aptos"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Aptos")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Aptos")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    node = props.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        props.append(node)
    node.set(qn("w:fill"), fill)


def page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    font(run, 8, color=MUTED)
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def add_title(doc, number, title, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    font(p.add_run(f"{number:02d}"), 10, True, ORANGE)
    font(p.add_run(f"  {title}"), 20, True, BLUE)
    if subtitle:
        q = doc.add_paragraph()
        q.paragraph_format.space_after = Pt(7)
        font(q.add_run(subtitle), 9.2, False, MUTED, True)


def add_paragraph(doc, text, lead=None, size=9.2, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.03
    p.paragraph_format.space_after = Pt(after)
    if lead:
        font(p.add_run(lead), size, True)
    font(p.add_run(text), size)
    return p


def add_bullets(doc, items, size=8.9):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.45)
        p.paragraph_format.first_line_indent = Cm(-0.2)
        p.paragraph_format.space_after = Pt(2)
        font(p.add_run(item), size)


def add_callout(doc, title, body, fill=PALE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    font(p.add_run(title + " "), 9.2, True, BLUE)
    font(p.add_run(body), 9.2)


def add_image(doc, filename, width=6.65):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    shape = p.add_run().add_picture(str(IMAGES / filename), width=Inches(width))
    shape._inline.docPr.set("descr", filename.replace("_", " "))


def add_caption(doc, label, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    font(p.add_run(label + ". "), 8.1, True, BLUE)
    font(p.add_run(text), 8.1, False, MUTED)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for cell, value in zip(table.rows[0].cells, headers):
        shade(cell, BLUE)
        p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
        font(p.add_run(value), 7.8, True, WHITE)
    for i, values in enumerate(rows):
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            shade(cell, WHITE if i % 2 == 0 else "F3F6F8")
            p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
            font(p.add_run(str(value)), 7.6)
    if widths:
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = Cm(width)
    return table


def next_page(doc):
    doc.add_page_break()


def configure() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7); section.page_width = Cm(21)
    section.top_margin = Cm(1.45); section.bottom_margin = Cm(1.35)
    section.left_margin = Cm(1.65); section.right_margin = Cm(1.65)
    section.header_distance = Cm(0.65); section.footer_distance = Cm(0.65)
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"; normal.font.size = Pt(9.2)
    normal.paragraph_format.space_after = Pt(4)
    header = section.header.paragraphs[0]
    font(header.add_run("FAKENEWS-ABM  |  Informe consolidado"), 7.8, True, MUTED)
    page_number(section.footer.paragraphs[0])
    return doc


def build():
    doc = configure()

    # 1. Cover and executive capsule.
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(32); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("FAKENEWS-ABM"), 14, True, ORANGE)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(8)
    font(p.add_run("Memoria, WOM y estrategias de camuflaje\nen la diseminación de noticias falsas"), 25, True, BLUE)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("Reporte consolidado de 2.903 simulaciones"), 12, True, INK)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(18)
    font(p.add_run("Paul Leger, Agustín Olivares, Oswaldo Teran, Manuela Lopez,\nFrancis Espinoza y Carolina Rodríguez"), 9.5, False, MUTED)
    add_callout(doc, "Resultado central.", "Una fuente puede volverse mucho más seleccionada sin aumentar la cantidad de noticias falsas. La diferencia depende de qué atributos copia, cuándo lo hace, cuánto recuerdan los usuarios y cómo WOM valora las noticias verdaderas y falsas.")
    add_paragraph(doc, "El estudio usa un modelo basado en agentes como laboratorio computacional. Cada usuario evalúa fuentes, selecciona una y puede transmitir su experiencia a contactos. Las intervenciones hipotéticas permiten observar mecanismos que serían difíciles o poco éticos de manipular en plataformas reales.", lead="Objetivo. ", size=10, after=8)
    add_table(doc, ["Fase", "Propósito", "Simulaciones"], [
        ["Exploratoria", "Memoria, momento, alcance de copia, WOM y semillas", "814"],
        ["Baseline", "Comprobar tasas y WOM sensible a la veracidad", "178"],
        ["Tratamientos", "Estrategias, confirmación y robustez", "1.911"],
        ["Total", "Evidencia consolidada", "2.903"],
    ], [3.0, 10.3, 3.0])
    add_paragraph(doc, "Proyecto PLURALISMO PLU230018 · Universidad Católica del Norte · Agosto de 2026", size=8.2, after=0)

    # 2. Executive summary.
    next_page(doc); add_title(doc, 2, "Resumen ejecutivo", "Qué aprendimos y por qué importa")
    add_paragraph(doc, "El modelo reproduce una ecología con medios tradicionales, medios desconocidos, una fuente especializada en fake news y una fuente mixta. El resultado primario no es solo qué fuente se vuelve popular, sino qué proporción de las decisiones de repost corresponde efectivamente a una noticia falsa.")
    add_bullets(doc, [
        "El baseline pasó cinco de cinco comprobaciones. Las tasas simuladas de publicación falsa quedaron próximas a las probabilidades configuradas.",
        "WOM sensible a la veracidad redujo los reposts falsos de los últimos 100 períodos en 5,26 puntos porcentuales frente a WOM que solo permite descubrimiento.",
        "Copiar todos los atributos puede aumentar la selección de la fuente falsa hasta 8,54 puntos, pero disminuir los reposts falsos porque también copia la credibilidad y, con ella, una probabilidad mucho menor de mentir.",
        "Copiar solamente atributos de engagement o no relacionados con credibilidad es la estrategia más limpia para estudiar camuflaje: conserva la alta propensión de la fuente a publicar falsedades.",
        "El mayor aumento confirmado de diseminación fue 2,89 puntos porcentuales, con intervención en el período 1, memoria 25, penalización de recomendaciones falsas e indiferencia ante recomendaciones verdaderas.",
        "Las intervenciones tempranas dominan a las tardías. Con memorias muy cortas, el efecto prácticamente desaparece.",
        "Las semillas comunes reprodujeron el orden de condiciones y, en ocho de diez comparaciones, redujeron la incertidumbre.",
    ], 9.1)
    add_callout(doc, "Lectura prudente.", "Estos valores son efectos dentro del mundo simulado. Sirven para comparar mecanismos y formular hipótesis; no son estimaciones directas del comportamiento de una plataforma real.")
    add_paragraph(doc, "Para periodistas y especialistas en comunicación, el resultado distingue visibilidad de desinformación. Para marketing, muestra que rasgos atractivos pueden transferir audiencia. Para ciencias sociales, evidencia dependencias temporales e interacción social. Para ingeniería, demuestra la utilidad de diseños reproducibles y emparejados.")

    # 3. Model.
    next_page(doc); add_title(doc, 3, "Cómo funciona el modelo", "Una explicación sencilla de agentes, endosos, memoria y WOM")
    add_paragraph(doc, "Un agente representa a una persona usuaria de una red social. En cada período conoce una o más fuentes, evalúa sus características y selecciona una noticia para repostear. Una selección es, por tanto, una decisión observable dentro del modelo.", lead="Usuarios y fuentes. ")
    add_paragraph(doc, "Los endosos son pequeñas piezas de evidencia a favor o en contra de una fuente. Se construyen al observar atributos como emoción, lenguaje simple, cercanía, sensacionalismo, calidad, entretenimiento, credibilidad, material audiovisual, hashtags y enlaces.", lead="Endosos. ")
    add_paragraph(doc, "La memoria indica cuántos períodos de evidencia siguen influyendo. Memoria 5 significa que la información antigua desaparece rápidamente; memoria 25 o 100 conserva una historia más larga; memoria infinita significa que la evidencia acumulada nunca expira.", lead="Memoria. ")
    add_paragraph(doc, "WOM —word of mouth o boca a boca— permite que una persona recomiende una fuente a sus contactos. La recomendación puede penalizarse, ignorarse o premiarse según la noticia recomendada resulte falsa o verdadera.", lead="WOM. ")
    add_table(doc, ["Política", "Valor", "Interpretación"], [
        ["Penalizar", "-1", "La experiencia reduce la valoración transmitida"],
        ["Ignorar", "0", "No se crea un endoso por ese resultado"],
        ["Premiar", "+1", "La experiencia aumenta la valoración transmitida"],
    ], [3.1, 2.1, 11.2])
    add_callout(doc, "Punto clave.", "WOM=1 habilita recomendaciones. Los valores -1, 0 y +1 no cambian el volumen de WOM: definen cómo se valora el resultado conocido de la noticia.")
    add_paragraph(doc, "La credibilidad cumple dos funciones en la implementación actual: participa en la evaluación de la fuente y controla su probabilidad de publicar fake news. Esta doble función obliga a interpretar con cuidado los escenarios que copian credibilidad.")

    # 4. Method.
    next_page(doc); add_title(doc, 4, "Metodología y calidad de la evidencia", "Diseño factorial, controles emparejados y métricas")
    add_paragraph(doc, "Cada condición principal utiliza 400 agentes, cuatro fuentes y 400 períodos. La ventana primaria promedia los períodos 301 a 400, cuando la simulación ya contiene una historia sustancial. Las intervenciones se comparan con controles de igual configuración y semilla.")
    add_table(doc, ["Métrica", "Definición", "Pregunta que responde"], [
        ["Selección objetivo", "Selecciones de FAKE_NEWS_SOURCE / 400", "¿La fuente ganó audiencia?"],
        ["Reposts falsos", "Selecciones ponderadas por estado falso / 400", "¿Circuló más desinformación?"],
        ["Acumulado falso", "Suma de reposts falsos en 400 períodos", "¿Cuál fue la carga total?"],
        ["Tasa falsa objetivo", "Períodos falsos / 400", "¿Cambió la conducta editorial?"],
    ], [3.1, 7.2, 6.1])
    add_paragraph(doc, "Una semilla común intenta exponer tratamiento y control a la misma secuencia aleatoria inicial. El efecto se calcula como tratamiento menos control para cada semilla y luego se promedian las diferencias. Los intervalos del 95% usan la distribución t de Student sobre estas diferencias.", lead="Emparejamiento. ")
    add_bullets(doc, [
        "178 corridas para baseline y sensibilidad de memoria.",
        "1.287 corridas para la matriz principal de estrategias.",
        "360 corridas confirmatorias con 30 semillas por política WOM.",
        "264 corridas de robustez variando contactos y alcance.",
        "814 corridas exploratorias anteriores, reportadas como fase separada.",
    ])
    add_callout(doc, "Validaciones estructurales.", "Cada corrida contiene 400 filas de períodos, estados falsos binarios, claves de período únicas y totales de selección compatibles con la configuración de alcance.")
    add_paragraph(doc, "No se aplicó corrección por comparaciones múltiples. Por ello, el ranking amplio es exploratorio; las seis condiciones confirmatorias, con 30 semillas, reciben mayor peso interpretativo.")

    # 5. Scenarios.
    next_page(doc); add_title(doc, 5, "Escenarios estudiados", "Qué cambia en cada intervención hipotética")
    add_table(doc, ["Escenario", "Qué copia", "Uso analítico"], [
        ["Sin escenario", "Nada", "Control comparable"],
        ["Credibility", "Solo credibilidad", "Diagnóstico de probabilidad falsa"],
        ["Non-credibility", "Todos menos credibilidad", "Camuflaje principal"],
        ["All", "Todos los atributos", "Máxima imitación, pero cambia veracidad"],
    ], [3.2, 5.2, 8.0])
    add_paragraph(doc, "La fuente objetivo es FAKE_NEWS_SOURCE y la fuente donante es un medio tradicional. La intervención puede comenzar en los períodos 1, 25, 50 o 100. Se estudian memorias 5, 10, 25, 50, 100 e infinita en la fase exploratoria; la matriz nueva concentra la comparación en 25, 100 e infinita.")
    add_paragraph(doc, "La matriz de políticas combina tres respuestas ante una noticia falsa (-1, 0 o +1) y dos respuestas ante una noticia verdadera (0 o +1). Esto separa cuatro ideas: castigar o ignorar falsedades, y premiar o ignorar verdades.", lead="Políticas WOM. ")
    add_paragraph(doc, "La robustez varía 6, 15 y 27 contactos potenciales, junto con alcance desactivado o probabilidades objetivo de 14,7%, 46,7% y 81,2%. El modelo no tiene un parámetro de homofilia; por tanto, no puede responder todavía si usuarios similares se conectan preferentemente.", lead="Red y alcance. ")
    add_callout(doc, "Dos preguntas distintas.", "All-copy pregunta cuánta audiencia gana una copia total. Non-credibility pregunta cuánta desinformación puede producir una fuente que mantiene su conducta falsa pero adopta una presentación más competitiva.")
    add_paragraph(doc, "La fase de calibración empírica quedó preparada, pero no ejecutada: el repositorio no contiene datos observados de exposición, repost, credibilidad ni incertidumbre muestral. Ajustar parámetros sin esos datos produciría una precisión ficticia.")

    # 6. Baseline.
    next_page(doc); add_title(doc, 6, "Resultado 1: el baseline se comporta como se esperaba", "WOM sensible a la veracidad reduce la circulación falsa")
    add_image(doc, "01_baseline_wom.png", 6.55)
    add_caption(doc, "Figura 1", "Promedio de reposts falsos en los períodos 301-400. B0 y B1 coinciden; B2 incorpora penalización de falsedades y premio de verdades.")
    add_paragraph(doc, "No WOM y WOM de descubrimiento produjeron exactamente 15,36% de reposts falsos en esta comparación emparejada. Esto indica que permitir recomendaciones sin valorarlas por su resultado no cambió la decisión agregada bajo esta configuración.")
    add_paragraph(doc, "Al penalizar recomendaciones falsas y premiar las verdaderas, la proporción bajó a 10,11%. La diferencia fue -5,26 puntos porcentuales, con IC 95% [-5,95; -4,56] y 30 pares.", lead="Efecto principal. ")
    add_table(doc, ["Comprobación", "Esperado", "Observado"], [
        ["Medio tradicional: tasa falsa", "9,2%", "9,26%"],
        ["Medio desconocido: tasa falsa", "20,6%", "20,58%"],
        ["Fuente fake news: tasa falsa", "66,7%", "67,24%"],
        ["Fuente mixta: tasa falsa", "41,6%", "41,64%"],
        ["Efecto WOM sensible", "Negativo", "-5,26 pp"],
    ], [7.3, 4.3, 4.8])
    add_callout(doc, "Interpretación.", "El baseline no demuestra que WOM siempre sea protector. Demuestra que, cuando la recomendación incorpora correctamente la veracidad observada, el mecanismo implementado puede reducir la desinformación.")

    # 7. Memory.
    next_page(doc); add_title(doc, 7, "Resultado 2: memoria y momento interactúan", "La intervención temprana necesita evidencia que permanezca")
    add_image(doc, "02_memoria_momento.png", 6.55)
    add_caption(doc, "Figura 2", "Aumento de selección de la fuente transformada frente al control. Las celdas muestran puntos porcentuales; 'Infinite' conserva toda la evidencia.")
    add_paragraph(doc, "Con memoria 5, la intervención casi no deja huella: la evidencia asociada a los atributos copiados expira antes de sostener una ventaja. Con memoria 10 aparece un efecto pequeño y concentrado al comienzo.")
    add_paragraph(doc, "Desde memoria 25, activar en el período 1 es sistemáticamente más influyente que hacerlo en 25, 50 o 100. La fase exploratoria alcanzó 8,34 puntos de aumento con memoria 100 y período 1; memoria infinita fue similar, con 8,20 puntos.", lead="Patrón temporal. ")
    add_paragraph(doc, "La matriz nueva de camuflaje confirma la caída con el retraso: para memoria infinita y política -1/+1, el efecto sobre reposts falsos baja desde 1,76 puntos en período 1 a 0,66 en período 25 y continúa debilitándose.")
    add_callout(doc, "Significado social.", "La memoria no debe entenderse como una capacidad psicológica literal de 25 días. Es el horizonte durante el cual la experiencia conserva influencia en el mecanismo de decisión.")

    # 8. Copy scope.
    next_page(doc); add_title(doc, 8, "Resultado 3: copiar más no siempre disemina más", "Popularidad, credibilidad y conducta editorial son dimensiones diferentes")
    add_image(doc, "03_alcance_copia.png", 6.45)
    add_caption(doc, "Figura 3", "Fase exploratoria con memoria infinita y WOM activo. La copia total eleva la selección, pero reduce drásticamente la tasa falsa de la fuente.")
    add_paragraph(doc, "Sin escenario, la fuente objetivo fue seleccionada cerca de 0,81% y publicó falsedades en 67,36% de los períodos. Al copiar solo atributos de engagement, la selección aumentó a 4,29% y la tasa falsa permaneció en 65,82%.")
    add_paragraph(doc, "Con copia total, la selección subió a 9,22%, pero la tasa falsa cayó a 9,82%. En la matriz nueva, la combinación all-copy, período 1 y memoria infinita aumentó la selección en 8,54 puntos, al mismo tiempo que redujo los reposts falsos en 1,16 puntos.", lead="Paradoja aparente. ")
    add_paragraph(doc, "No es una contradicción: la fuente gana audiencia, pero se comporta como un medio tradicional porque copió credibilidad. Por eso, all-copy es útil para estudiar imitación completa, pero no representa una estrategia en la que el emisor conserva su intención de desinformar.")
    add_callout(doc, "Conclusión operacional.", "Para estudiar diseminación maliciosa, la métrica primaria debe ser el repost falso observado y el tratamiento principal debe conservar la probabilidad original de publicar falsedades.")

    # 9. WOM policies.
    next_page(doc); add_title(doc, 9, "Resultado 4: la política WOM cambia el efecto del camuflaje", "Ignorar, penalizar o premiar no son equivalentes")
    add_image(doc, "05_politicas_wom.png", 6.45)
    add_caption(doc, "Figura 4", "Efecto confirmatorio del camuflaje non-credibility activado en período 1, memoria 25 y 30 semillas. Barras: puntos porcentuales; líneas: IC 95%.")
    add_table(doc, ["Falsa / verdadera", "Reposts falsos", "Selección objetivo", "Acumulado"], [
        ["-1 / 0", "+2,89 pp", "+5,64 pp", "+4.354"],
        ["0 / 0", "+1,36 pp", "+2,65 pp", "+2.698"],
        ["+1 / 0", "+0,57 pp", "+1,10 pp", "+2.083"],
        ["-1 / +1", "+0,02 pp", "+0,04 pp", "+859"],
        ["0 / +1", "+0,02 pp", "+0,03 pp", "+797"],
        ["+1 / +1", "≈0,00 pp", "+0,02 pp", "+766"],
    ], [4.1, 4.0, 4.0, 3.8])
    add_paragraph(doc, "El mayor efecto relativo aparece al penalizar noticias falsas pero no premiar noticias verdaderas. Esto no significa que penalizar falsedades sea dañino en términos absolutos: el contraste mide cuánto añade el camuflaje frente a un control sometido a la misma política.")
    add_callout(doc, "No confundir estimandos.", "La política -1/+1 protege el baseline. La tabla pregunta algo distinto: cuánto cambia una intervención de camuflaje dentro de cada política.")

    # 10. Networks.
    next_page(doc); add_title(doc, 10, "Resultado 5: red, alcance y semillas", "La robustez depende de la forma de comparación")
    add_image(doc, "06_robustez_red.png", 6.35)
    add_caption(doc, "Figura 5", "Efecto de camuflaje sobre reposts falsos bajo 6, 15 o 27 contactos y cuatro configuraciones de alcance. La mayoría de los efectos es pequeña.")
    add_paragraph(doc, "En la fase exploratoria, comparar WOM activo con WOM inactivo redujo la selección de la fuente transformada en las seis combinaciones soportadas de grado y alcance. Los efectos variaron entre -1,42 y -2,67 puntos; los reposts falsos también disminuyeron.")
    add_paragraph(doc, "La prueba nueva pregunta cuánto añade el camuflaje cuando WOM ya está activo y sensible a la veracidad. Allí los efectos son generalmente positivos, pero pequeños: de aproximadamente 0,00 a 0,18 puntos. Las dos afirmaciones pueden coexistir porque sus controles son diferentes.")
    add_image(doc, "07_precision_semillas.png", 5.55)
    add_caption(doc, "Figura 6", "Razón entre error estándar emparejado y no emparejado. Valores bajo 1 representan mayor precisión con semillas comunes.")
    add_paragraph(doc, "El orden de 20 condiciones se reprodujo con correlaciones de Spearman 0,998 para selección y 0,974 para reposts falsos. El emparejamiento redujo incertidumbre en ocho de diez celdas.")

    # 11. Most promoting combinations.
    next_page(doc); add_title(doc, 11, "Qué combinaciones promueven más la fuente falsa", "Ranking separado para popularidad y desinformación")
    add_image(doc, "04_popularidad_vs_desinformacion.png", 6.25)
    add_caption(doc, "Figura 7", "Cada punto es una estrategia. Derecha significa más selección; arriba significa más reposts falsos. All-copy se concentra a la derecha y abajo; non-credibility, arriba y a la derecha.")
    add_table(doc, ["Objetivo", "Combinación más fuerte", "Efecto"], [
        ["Aumentar selección", "All-copy · P1 · memoria infinita · -1/+1", "+8,54 pp"],
        ["Aumentar repost falso", "Non-credibility · P1 · memoria 25 · -1/0", "+2,89 pp"],
        ["Aumentar carga acumulada", "Misma condición confirmatoria", "+4.354 reposts"],
        ["Reducir baseline falso", "WOM sensible -1/+1", "-5,26 pp"],
    ], [4.0, 8.5, 3.6])
    add_paragraph(doc, "Para maximizar audiencia simulada, la combinación dominante imita todos los atributos, actúa desde el inicio y conserva evidencia durante mucho tiempo. Sin embargo, deja de ser una fuente altamente falsa debido a la copia de credibilidad.")
    add_paragraph(doc, "Para maximizar desinformación sin cambiar la propensión editorial, la combinación más preocupante copia los atributos de presentación, actúa desde el inicio y usa una memoria suficientemente larga. En la confirmación, no premiar noticias verdaderas amplifica fuertemente la diferencia entre tratamiento y control.")
    add_callout(doc, "Regla de lectura.", "Nunca usar participación de mercado como sustituto automático de desinformación. Una estrategia debe evaluarse simultáneamente por audiencia, tasa de publicación falsa y reposts falsos efectivos.")

    # 12. Discussion and references.
    next_page(doc); add_title(doc, 12, "Discusión, implicaciones y conclusiones", "Qué significa la evidencia y qué falta por aprender")
    add_paragraph(doc, "El estudio muestra que la diseminación emerge de una interacción entre atractivo, credibilidad, memoria, momento e intercambio social. No existe una única variable que determine el resultado. La misma intervención puede ganar audiencia y reducir falsedades, o conservar una audiencia modesta pero aumentar la carga falsa.")
    add_bullets(doc, [
        "Periodismo: evaluar tácticas de presentación junto con conducta editorial; la apariencia tradicional no garantiza veracidad.",
        "Marketing y comunicación: engagement aumenta descubrimiento, pero sus efectos dependen del historial y de la respuesta social a resultados verdaderos y falsos.",
        "Ciencias sociales: el horizonte de memoria y la secuencia de intervención son mecanismos sustantivos, no detalles técnicos.",
        "Ingeniería de plataformas: conviene distinguir recomendación, valoración posterior y probabilidad de exposición, y registrar cada componente por separado.",
    ])
    add_paragraph(doc, "El modelo no está calibrado con datos empíricos, no implementa homofilia y representa fuentes como tipos agregados. Los intervalos cuantifican variación Monte Carlo, no incertidumbre sobre la sociedad. Algunas comparaciones exploratorias son numerosas y no tienen ajuste por multiplicidad.", lead="Limitaciones. ")
    add_paragraph(doc, "La siguiente etapa debe incorporar datos observados de exposición, repost y credibilidad; calibrar WOM, memoria y alcance; reservar datos para validación fuera de muestra; y luego incorporar homofilia. También es recomendable preregistrar una matriz reducida de hipótesis confirmatorias.", lead="Siguientes pasos. ")
    add_callout(doc, "Conclusión.", "El camuflaje más riesgoso no es copiar todo, sino copiar tempranamente lo que vuelve atractiva a una fuente mientras se conserva su alta probabilidad de publicar falsedades.")
    add_paragraph(doc, "Referencias seleccionadas", lead="", size=8.3, after=2)
    refs = [
        "Cohen, P. R. (1985). Heuristic Reasoning about Uncertainty: An Artificial Intelligence Approach.",
        "Grimm, V. et al. (2020). The ODD protocol for describing agent-based models. Journal of Artificial Societies and Social Simulation, 23(2).",
        "Lazer, D. M. J. et al. (2018). The science of fake news. Science, 359, 1094-1096.",
        "Terán, O., Leger, P., & López, M. (2022). Modeling and simulating Chinese cross-border e-commerce. Journal of Simulation. https://doi.org/10.1080/17477778.2022.2043791",
        "Terán, O., Leger, P., & López, M. (2024). Factors that drive market share and the oligopolistic character of cross-border B2C e-commerce. Simulation. https://doi.org/10.1177/00375497241296542",
        "Vosoughi, S., Roy, D., & Aral, S. (2018). The spread of true and false news online. Science, 359, 1146-1151.",
    ]
    for ref in refs:
        add_paragraph(doc, ref, size=7.4, after=1)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
