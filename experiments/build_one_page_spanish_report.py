#!/usr/bin/env python3
"""Genera un informe ejecutivo de seis páginas a partir de resultados validados."""

from __future__ import annotations

import argparse
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from matplotlib.ticker import FuncFormatter
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLUE = "1F4E79"
TEAL = "087F8C"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F2F4F7"
DARK = "17212B"
MUTED = "5D6874"
WHITE = "FFFFFF"
GOLD = "FFF3CD"
TABLE_WIDTH = 9360
TABLE_INDENT = 120


def set_font(run, size: float, color: str = DARK, bold: bool = False) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def set_cell_fill(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, dxa: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    width = properties.find(qn("w:tcW"))
    if width is None:
        width = OxmlElement("w:tcW")
        properties.append(width)
    width.set(qn("w:w"), str(dxa))
    width.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    properties = table._tbl.tblPr
    width = properties.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:w"), str(sum(widths)))
    width.set(qn("w:type"), "dxa")
    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), str(TABLE_INDENT))
    indent.set(qn("w:type"), "dxa")
    layout = properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)
    for row in table.rows:
        for cell, value in zip(row.cells, widths):
            set_cell_width(cell, value)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    properties.append(marker)


def add_heading(document, text: str, level: int = 1, before: float | None = None) -> None:
    paragraph = document.add_paragraph(style=f"Heading {level}")
    if before is not None:
        paragraph.paragraph_format.space_before = Pt(before)
    paragraph.add_run(text)


def add_lead_paragraph(document, lead: str, body: str, after=4, size=9.5) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.08
    set_font(paragraph.add_run(lead), size, DARK, True)
    set_font(paragraph.add_run(body), size, DARK)


def add_callout(document, lead: str, body: str, fill=LIGHT_BLUE) -> None:
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [TABLE_WIDTH])
    cell = table.cell(0, 0)
    set_cell_fill(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.08
    set_font(paragraph.add_run(lead), 9.2, DARK, True)
    set_font(paragraph.add_run(body), 9.2, DARK)
    set_repeat_header(table.rows[0])


def set_picture_alt_text(inline_shape, description: str) -> None:
    inline_shape._inline.docPr.set("descr", description)
    inline_shape._inline.docPr.set("title", description)


def add_data_table(document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_fill(cell, BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        set_font(paragraph.add_run(text), 8.6, WHITE, True)
    set_repeat_header(table.rows[0])
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for column, (cell, value) in enumerate(zip(cells, values)):
            set_cell_fill(cell, WHITE if row_index % 2 == 0 else LIGHT_GRAY)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.02
            if column > 0:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(paragraph.add_run(value), 8.3, DARK, column == 0)
    set_table_geometry(table, widths)


def make_memory_chart(q2: pd.DataFrame, output: Path) -> None:
    memories = [5, 10, 25, 50, 100, -1]
    periods = [1, 25, 50, 100]
    pivot = 100 * q2.pivot(index="memory", columns="timing_period", values="target_share_effect_vs_control")
    matrix = pivot.loc[memories, periods].to_numpy()
    fig, axis = plt.subplots(figsize=(7.2, 3.25), dpi=180)
    image = axis.imshow(matrix, cmap="Blues", vmin=0, vmax=8.5, aspect="auto")
    for row in range(matrix.shape[0]):
        for column in range(matrix.shape[1]):
            value = matrix[row, column]
            axis.text(column, row, f"{value:.2f}".replace(".", ","), ha="center", va="center",
                      fontsize=9, color=f"#{WHITE}" if value > 5 else f"#{DARK}")
    axis.set_xticks(range(len(periods)), [str(value) for value in periods])
    axis.set_yticks(range(len(memories)), ["Sin límite" if value == -1 else str(value) for value in memories])
    axis.set_xlabel("Período de activación del escenario", fontsize=9)
    axis.set_ylabel("Memoria (períodos)", fontsize=9)
    axis.set_title("Aumento de participación de la fuente transformada", loc="left", fontsize=11, weight="bold")
    colorbar = fig.colorbar(image, ax=axis, fraction=0.035, pad=0.03)
    colorbar.set_label("Puntos porcentuales vs. sin escenario", fontsize=8)
    colorbar.ax.tick_params(labelsize=8)
    axis.tick_params(labelsize=8.5)
    for spine in axis.spines.values():
        spine.set_color("#D7DEE5")
    fig.tight_layout()
    fig.savefig(output, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def make_wom_chart(q3: pd.DataFrame, output: Path) -> None:
    ordered = q3.sort_values(["source_reach", "realized_degree"])
    labels = [f"Grado {int(row.realized_degree)} | alcance {'sí' if row.source_reach else 'no'}"
              for row in ordered.itertuples()]
    values = 100 * ordered.target_share_wom_effect.to_numpy()
    low = 100 * ordered.target_share_wom_effect_ci95_low.to_numpy()
    high = 100 * ordered.target_share_wom_effect_ci95_high.to_numpy()
    errors = np.vstack([values - low, high - values])
    fig, axis = plt.subplots(figsize=(7.2, 3.05), dpi=180)
    positions = np.arange(len(labels))
    axis.barh(positions, values, color="#4C78A8", edgecolor="#274B6D", height=0.62)
    axis.errorbar(values, positions, xerr=errors, fmt="none", ecolor=f"#{DARK}", capsize=3, linewidth=1)
    axis.axvline(0, color="#56616C", linewidth=1)
    for y, value in zip(positions, values):
        axis.text(-0.08, y, f"{value:.2f}".replace(".", ","), va="center", ha="right", color=f"#{WHITE}", fontsize=8.5, weight="bold")
    axis.set_yticks(positions, labels)
    axis.invert_yaxis()
    axis.set_xlim(-3.55, 0.15)
    axis.xaxis.set_major_formatter(FuncFormatter(lambda value, _: f"{value:.1f}".replace(".", ",")))
    axis.set_xlabel("Efecto de WOM sobre la participación (puntos porcentuales)", fontsize=8.5)
    axis.set_title("Efecto del boca a boca (WOM) según red y alcance", loc="left", fontsize=11, weight="bold")
    axis.grid(axis="x", color="#E3E7EB", linewidth=0.7)
    axis.set_axisbelow(True)
    axis.tick_params(labelsize=8.3)
    for side in ("top", "right", "left"):
        axis.spines[side].set_visible(False)
    fig.tight_layout()
    fig.savefig(output, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def make_selection_ranking_chart(q1: pd.DataFrame, output: Path) -> pd.DataFrame:
    """Compara las combinaciones Q1 más informativas para seleccionar la fuente objetivo."""

    def pick(scope: str, memory: int, wom: int) -> pd.Series:
        row = q1[(q1.scope == scope) & (q1.memory == memory) & (q1.wom == wom)]
        if len(row) != 1:
            raise ValueError(f"Expected one Q1 row for scope={scope}, memory={memory}, wom={wom}")
        return row.iloc[0]

    specifications = [
        ("Todos | M25 | WOM no | P1", "Todos", pick("all", 25, 0)),
        ("Todos | M∞ | WOM no | P1", "Todos", pick("all", -1, 0)),
        ("Todos | M∞ | WOM sí | P1", "Todos", pick("all", -1, 1)),
        ("Engagement | M∞ | WOM no | P1", "Engagement", pick("engagement", -1, 0)),
        ("Engagement | M∞ | WOM sí | P1", "Engagement", pick("engagement", -1, 1)),
        ("Engagement | M25 | WOM no | P1", "Engagement", pick("engagement", 25, 0)),
        ("Sin escenario | M∞ | WOM no", "Sin escenario", pick("none", -1, 0)),
        ("Engagement | M25 | WOM sí | P1", "Engagement", pick("engagement", 25, 1)),
    ]
    ranking = pd.DataFrame(
        {
            "label": [item[0] for item in specifications],
            "family": [item[1] for item in specifications],
            "share": [item[2].target_share_final100_mean for item in specifications],
            "fake_rate": [item[2].target_fake_rate_final100_mean for item in specifications],
        }
    )
    ranking["users"] = 400 * ranking.share
    ranking = ranking.sort_values("share", ascending=True).reset_index(drop=True)

    colors = {
        "Todos": "#2C5F8A",
        "Engagement": "#D4872C",
        "Sin escenario": "#8A949E",
    }
    fig, axis = plt.subplots(figsize=(7.2, 3.5), dpi=180)
    positions = np.arange(len(ranking))
    bars = axis.barh(
        positions,
        100 * ranking.share,
        color=[colors[value] for value in ranking.family],
        edgecolor="#263746",
        height=0.62,
    )
    for bar, share, users in zip(bars, 100 * ranking.share, ranking.users):
        user_label = "<1 usuario" if users < 0.5 else f"{users:.0f} usuarios"
        label = f"{share:.2f}% | {user_label}".replace(".", ",")
        axis.text(share + 0.18, bar.get_y() + bar.get_height() / 2, label,
                  va="center", ha="left", fontsize=8.2, color=f"#{DARK}")
    axis.set_yticks(positions, ranking.label)
    axis.set_xlim(0, 16.2)
    axis.xaxis.set_major_formatter(FuncFormatter(lambda value, _: f"{value:.0f}%"))
    axis.set_xlabel("Participación de FAKE_NEWS_SOURCE en los períodos 301-400", fontsize=8.5)
    axis.set_title("Selección de la fuente objetivo por combinación", loc="left", fontsize=11, weight="bold")
    axis.grid(axis="x", color="#E3E7EB", linewidth=0.7)
    axis.set_axisbelow(True)
    axis.tick_params(labelsize=8.1)
    for side in ("top", "right", "left"):
        axis.spines[side].set_visible(False)
    fig.tight_layout()
    fig.savefig(output, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return ranking.sort_values("share", ascending=False).reset_index(drop=True)


def configure_document() -> Document:
    document = Document()
    section = document.sections[0]
    # Named override: two_page_research_brief. It preserves the business-brief
    # hierarchy while using a compact vertical page budget for two full pages.
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.68)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.10

    for level, size, before, after in ((1, 15, 10, 5), (2, 12.5, 8, 4)):
        style = document.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(BLUE)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("FAKENEWS-ABM | Síntesis de experimentos"), 8, MUTED)
    return document


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("experiment_root", type=Path)
    parser.add_argument("--output", type=Path)
    args = parser.parse_args()
    analysis = args.experiment_root / "analysis"
    output = args.output or analysis / "reporte_experimentos_es_6_paginas.docx"

    q1 = pd.read_csv(analysis / "q1_copy_scope.csv")
    q2 = pd.read_csv(analysis / "q2_memory_timing_effects.csv")
    q3 = pd.read_csv(analysis / "q3_wom_network_effects.csv")
    rank = pd.read_csv(analysis / "q5_rank_reproduction.csv")
    precision = pd.read_csv(analysis / "q5_paired_precision.csv")
    runs = pd.read_csv(analysis / "run_metrics.csv")

    figures = analysis / "figures"
    figures.mkdir(parents=True, exist_ok=True)
    memory_chart = figures / "reporte_es_memoria_escenario.png"
    wom_chart = figures / "reporte_es_wom_red.png"
    selection_chart = figures / "reporte_es_ranking_seleccion.png"
    make_memory_chart(q2, memory_chart)
    make_wom_chart(q3, wom_chart)
    selection_ranking = make_selection_ranking_chart(q1, selection_chart)

    scope = q1.groupby("scope").agg(
        target=("target_share_final100_mean", "mean"),
        fake_rate=("target_fake_rate_final100_mean", "mean"),
    )
    target_rho = rank.loc[rank.metric == "target_share_final100", "spearman_rho"].iloc[0]
    fake_rho = rank.loc[rank.metric == "fake_repost_share_final100", "spearman_rho"].iloc[0]
    narrower = int((precision.se_ratio_paired_to_unpaired < 1).sum())

    document = configure_document()

    # Página 1: motivación, objetivo y resumen ejecutivo.
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(1)
    set_font(title.add_run("Difusión de noticias falsas: qué aprendimos"), 22, BLUE, True)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(8)
    set_font(subtitle.add_run("Síntesis ampliada de los experimentos realizados con FAKENEWS-ABM"), 10.5, MUTED)

    add_heading(document, "Objetivo del estudio", 1, before=0)
    add_lead_paragraph(document, "El estudio busca comprender bajo qué condiciones una fuente asociada a noticias falsas logra ser seleccionada y repostada por más usuarios. ",
                       "FAKENEWS-ABM funciona como un laboratorio computacional: representa usuarios de una red social, distintas fuentes informativas y decisiones repetidas de selección. Permite cambiar una condición a la vez (por ejemplo, memoria, WOM o momento del escenario) y observar cómo cambia la difusión.", after=5)
    add_lead_paragraph(document, "La pregunta no es solamente cuánto circula una fuente. ",
                       "También interesa distinguir si la fuente sigue publicando contenido falso. Una fuente puede ganar selecciones porque adopta atributos atractivos, pero puede dejar de comportarse como fuente falsa si el escenario copia también su probabilidad de publicar fake news.", after=5)

    add_heading(document, "Resumen ejecutivo", 1, before=5)
    add_lead_paragraph(document, "El escenario puede cambiar de manera sustantiva la competencia entre fuentes. ",
                       "La mayor selección absoluta aparece con memoria 25, WOM desactivado, escenario desde el período 1 y copia de todos los atributos: 13,99% de las decisiones, equivalentes a unos 56 de 400 usuarios por período.", after=3)
    add_lead_paragraph(document, "La memoria y el momento de intervención actúan conjuntamente. ",
                       "Los escenarios tempranos dejan una huella mayor cuando los agentes conservan suficiente historial. Con memoria 5, el efecto casi desaparece; con memoria 100 o ilimitada, una intervención en el período 1 agrega cerca de ocho puntos porcentuales frente a no intervenir.", after=3)
    add_lead_paragraph(document, "WOM no fue un amplificador automático. ",
                       "En todas las redes estudiadas, activar el boca a boca redujo la selección de la fuente transformada. También redujo la proporción total de reposts falsos, por lo que su efecto fue desfavorable para esa fuente, pero favorable para el sistema en términos de desinformación.", after=3)
    add_lead_paragraph(document, "La combinación que más promueve una fuente que sigue siendo falsa es distinta. ",
                       "Al copiar solo atributos de engagement, usar memoria ilimitada, desactivar WOM y comenzar en el período 1, la fuente obtiene 6,03% de selecciones y mantiene una tasa de publicaciones falsas de 65,9%.", after=5)

    add_heading(document, "Conceptos para leer el informe", 2, before=5)
    add_data_table(
        document,
        ["Concepto", "Explicación en palabras simples"],
        [
            ["Selección", "La fuente que un usuario elige en un período; en el modelo se interpreta como un repost."],
            ["Memoria", "Cantidad de períodos anteriores que el usuario recuerda al evaluar una fuente. MEMORY=-1 conserva todo el historial."],
            ["WOM", "Boca a boca: recomendaciones recibidas desde contactos de la red social."],
            ["Escenario", "Intervención que copia atributos desde una fuente de origen hacia una fuente objetivo en un período definido."],
        ],
        [2100, 7260],
    )

    document.add_page_break()

    # Página 2: memoria y momento del escenario.
    add_heading(document, "Memoria y momento: por qué intervenir temprano importa", 1, before=0)
    add_lead_paragraph(document, "En el modelo, memoria no significa capacidad psicológica general, sino una ventana temporal de experiencias. ",
                       "Con memoria 25, cada usuario evalúa las fuentes usando endorsements registrados durante los últimos 25 períodos. Los eventos más antiguos dejan de influir. Con memoria ilimitada (MEMORY=-1), todo el historial permanece disponible, incluidos los valores iniciales.", after=5)
    add_lead_paragraph(document, "Esta definición explica por qué el inicio del escenario importa. ",
                       "Una modificación aplicada temprano puede generar selecciones y experiencias posteriores que se acumulan. Si se aplica tarde, queda menos tiempo para alterar la trayectoria; si la memoria es muy corta, incluso una intervención temprana puede olvidarse rápidamente.", after=4)

    picture = document.add_paragraph()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.paragraph_format.space_after = Pt(1)
    memory_picture = picture.add_run().add_picture(str(memory_chart), width=Inches(6.0))
    set_picture_alt_text(memory_picture, "Matriz del aumento de participación de la fuente transformada según memoria y período de activación del escenario.")
    caption = document.add_paragraph()
    caption.paragraph_format.space_after = Pt(5)
    set_font(caption.add_run("Figura 1. "), 8, MUTED, True)
    set_font(caption.add_run("Cada celda muestra el cambio en puntos porcentuales frente al control sin escenario con la misma memoria. Valores más oscuros indican un aumento mayor de selección de la fuente objetivo."), 8, MUTED)

    add_heading(document, "Cómo leer la Figura 1", 2, before=3)
    add_lead_paragraph(document, "Las filas comparan cuánto recuerdan los agentes y las columnas cuándo se activa el escenario. ",
                       "La esquina inferior izquierda concentra los efectos mayores: memoria larga y activación en el período 1. La parte superior permanece casi blanca porque memorias 5 y 10 producen cambios pequeños, cualquiera sea el momento de aplicación.", after=3)
    add_lead_paragraph(document, "El contraste temporal es fuerte. ",
                       "Con memoria 100, comenzar en el período 1 eleva la participación en 8,34 puntos; comenzar en el período 100 la eleva solo 0,93 puntos. Con memoria 25, el mismo retraso reduce el efecto desde 2,06 hasta 0,05 puntos.", after=4)
    add_callout(document, "INTERPRETACIÓN. ", "El escenario no tiene un efecto fijo. Su resultado depende del historial que los agentes utilizan para decidir y del tiempo disponible para que la intervención reorganice ese historial.", fill=LIGHT_BLUE)

    document.add_page_break()

    # Página 3: alcance del escenario y significado de copiar atributos.
    add_heading(document, "Qué se copia cambia el significado del experimento", 1, before=0)
    add_lead_paragraph(document, "Se compararon tres condiciones. ",
                       "Sin escenario, la fuente conserva todos sus atributos originales. En la variante 'solo engagement', recibe atributos asociados al atractivo y la evaluación del contenido, pero conserva su probabilidad original de publicar noticias falsas. En la copia completa, recibe todos los atributos de la fuente de origen, incluida esa probabilidad.", after=5)

    add_data_table(
        document,
        ["Condición", "Participación de la fuente", "Usuarios esperados/400", "Publicaciones falsas"],
        [
            ["Sin escenario", f"{100 * scope.loc['none', 'target']:.2f}%".replace(".", ","), f"{400 * scope.loc['none', 'target']:.0f}", f"{100 * scope.loc['none', 'fake_rate']:.1f}%".replace(".", ",")],
            ["Solo atributos de engagement", f"{100 * scope.loc['engagement', 'target']:.2f}%".replace(".", ","), f"{400 * scope.loc['engagement', 'target']:.0f}", f"{100 * scope.loc['engagement', 'fake_rate']:.1f}%".replace(".", ",")],
            ["Todos los atributos", f"{100 * scope.loc['all', 'target']:.2f}%".replace(".", ","), f"{400 * scope.loc['all', 'target']:.0f}", f"{100 * scope.loc['all', 'fake_rate']:.1f}%".replace(".", ",")],
        ],
        [3300, 2050, 2050, 1960],
    )
    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(3)
    note.paragraph_format.space_after = Pt(5)
    set_font(note.add_run("Tabla 1. "), 8, MUTED, True)
    set_font(note.add_run("Promedios descriptivos sobre cuatro estratos de memoria y WOM. La columna 'usuarios esperados' traduce el porcentaje a una población de 400; no representa siempre un conteo entero observado en cada período."), 8, MUTED)

    add_heading(document, "Cómo interpretar la Tabla 1", 2, before=4)
    add_lead_paragraph(document, "Copiar solo engagement mejora la capacidad competitiva de la fuente sin volverla más confiable. ",
                       "La participación promedio sube de 0,47% a 3,27%, aproximadamente de 2 a 13 usuarios por período. Sin embargo, la tasa de publicaciones falsas permanece prácticamente igual: 67,1% sin escenario y 65,7% con copia de engagement.", after=4)
    add_lead_paragraph(document, "Copiar todos los atributos produce el mayor alcance promedio, pero cambia la identidad conductual de la fuente. ",
                       "La participación llega a 9,23%, unos 37 usuarios, mientras que la tasa de noticias falsas cae a 9,9%. Este caso muestra que una fuente originalmente etiquetada como FAKE_NEWS_SOURCE puede dejar de comportarse como tal si el escenario copia también la probabilidad de publicar contenido falso.", after=4)
    add_callout(document, "CONSECUENCIA METODOLÓGICA. ", "Para estudiar estrategias que hagan más atractiva una fuente falsa sin modificar su veracidad, la comparación apropiada es 'solo engagement'. La copia completa responde a otra pregunta: qué ocurre si la fuente adopta simultáneamente atractivo, credibilidad y propensión editorial de la fuente de origen.", fill=GOLD)

    add_heading(document, "Tres resultados que no deben confundirse", 2, before=5)
    add_lead_paragraph(document, "Participación de la fuente: ", "proporción de usuarios que selecciona FAKE_NEWS_SOURCE.  ", after=1, size=9.1)
    add_lead_paragraph(document, "Tasa de publicaciones falsas: ", "proporción de períodos en que esa fuente publica una noticia falsa.  ", after=1, size=9.1)
    add_lead_paragraph(document, "Reposts falsos del sistema: ", "selecciones ponderadas por el estado falso/verdadero de cada fuente; mide el resultado agregado de desinformación.", after=1, size=9.1)

    document.add_page_break()

    # Página 4: WOM y estructura de red.
    add_heading(document, "WOM no amplificó la fuente transformada", 1, before=0)
    add_lead_paragraph(document, "WOM representa el boca a boca entre contactos. ",
                       "Después de observar la selección de un contacto, un usuario puede recibir una recomendación que afectará la evaluación del período siguiente. La implementación asigna una contribución positiva si la noticia recomendada no es falsa y negativa si lo es.", after=5)
    add_lead_paragraph(document, "Por esta razón, WOM puede reforzar o penalizar una fuente. ",
                       "No debe asumirse que todo contacto social aumenta la difusión. En estos experimentos, su efecto neto fue negativo para la fuente transformada en todas las combinaciones de grado de contacto y alcance evaluadas.", after=3)

    picture = document.add_paragraph()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.paragraph_format.space_after = Pt(1)
    wom_picture = picture.add_run().add_picture(str(wom_chart), width=Inches(6.05))
    set_picture_alt_text(wom_picture, "Barras del efecto negativo de WOM sobre la participación de la fuente transformada en seis configuraciones de red y alcance.")
    caption = document.add_paragraph()
    caption.paragraph_format.space_after = Pt(5)
    set_font(caption.add_run("Figura 2. "), 8, MUTED, True)
    set_font(caption.add_run("Cada barra es WOM activado menos WOM desactivado. Un valor negativo significa que WOM reduce la selección de la fuente. Las líneas muestran intervalos aproximados de 95%."), 8, MUTED)

    add_heading(document, "Cómo leer la Figura 2", 2, before=3)
    add_lead_paragraph(document, "Las seis barras quedan a la izquierda de cero y sus intervalos también. ",
                       "La reducción varía entre 1,42 y 2,67 puntos porcentuales. El efecto más negativo aparece con grado 1 y alcance desactivado; el menos negativo, con grado 1 y alcance activado.", after=3)
    add_lead_paragraph(document, "El resultado sistémico apunta en la misma dirección. ",
                       "WOM redujo la proporción total de reposts falsos entre 1,61 y 6,19 puntos en las seis redes. Esto sugiere que, bajo la regla implementada, el contacto social transmite información que ayuda a penalizar contenido falso.", after=4)

    add_heading(document, "El efecto depende también de memoria y escenario", 2, before=3)
    add_data_table(
        document,
        ["Memoria y escenario", "Cambio por activar WOM", "Lectura"],
        [
            ["Ilimitada, sin escenario", "-0,21 puntos", "La fuente ya era minoritaria y pierde un poco más."],
            ["Ilimitada, escenario P1", "-1,77 puntos", "WOM reduce, pero no elimina, el impulso temprano."],
            ["25, sin escenario", "Aprox. 0,00 puntos", "La selección de la fuente ya es casi nula."],
            ["25, escenario P1", "-11,48 puntos", "WOM contrarresta gran parte de la ventaja creada por el escenario."],
        ],
        [3100, 2100, 4160],
    )
    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(3)
    set_font(note.add_run("Tabla 2. "), 8, MUTED, True)
    set_font(note.add_run("Contrastes emparejados WOM activado menos desactivado. El valor -11,48 no contradice la Figura 2: corresponde a otra matriz experimental, con copia completa y configuración de red por defecto."), 8, MUTED)

    document.add_page_break()

    # Página 5: ranking de combinaciones.
    add_heading(document, "Qué combinaciones promueven más la fuente de fake news", 1, before=0)
    add_lead_paragraph(document, "Esta sección ordena combinaciones representativas por selección de FAKE_NEWS_SOURCE durante los períodos 301-400. ",
                       "La lectura debe hacerse junto con la tasa de publicaciones falsas: el nombre de la fuente permanece igual, pero su comportamiento puede cambiar cuando se copian todos los atributos.", after=4)

    picture = document.add_paragraph()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.paragraph_format.space_after = Pt(1)
    selection_picture = picture.add_run().add_picture(str(selection_chart), width=Inches(6.1))
    set_picture_alt_text(selection_picture, "Ranking de ocho combinaciones por participación de FAKE_NEWS_SOURCE y número esperado de usuarios que la seleccionan por período.")
    caption = document.add_paragraph()
    caption.paragraph_format.space_after = Pt(5)
    set_font(caption.add_run("Figura 3. "), 8, MUTED, True)
    set_font(caption.add_run("Azul: copia de todos los atributos. Naranjo: solo engagement. Gris: sin escenario. M∞ indica memoria ilimitada y P1, activación en el período 1."), 8, MUTED)

    add_heading(document, "Cómo leer la Figura 3", 2, before=3)
    add_lead_paragraph(document, "Las tres barras superiores son copias completas. ",
                       "La mayor combina memoria 25, WOM desactivado y escenario en el período 1: 13,99%, unos 56 usuarios. Le siguen memoria ilimitada sin WOM (10,85%; 43 usuarios) y memoria ilimitada con WOM (9,22%; 37 usuarios). La diferencia entre las dos últimas vuelve a mostrar que WOM reduce la selección.", after=3)
    add_lead_paragraph(document, "Entre las condiciones que mantienen la fuente como predominantemente falsa, la mejor es la barra naranjo de 6,03%. ",
                       "Corresponde a copia de engagement, memoria ilimitada, WOM desactivado y activación en el período 1. Supera ampliamente el 1,06% del control sin escenario con la misma memoria y WOM.", after=4)

    ranking_rows = []
    for row in selection_ranking.head(5).itertuples():
        ranking_rows.append([row.label, f"{100 * row.share:.2f}%".replace(".", ","), f"{row.users:.0f}", f"{100 * row.fake_rate:.1f}%".replace(".", ",")])
    add_data_table(document, ["Combinación", "Selección", "Usuarios/400", "Publica fake news"], ranking_rows, [4440, 1400, 1500, 2020])
    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(3)
    note.paragraph_format.space_after = Pt(4)
    set_font(note.add_run("Tabla 3. "), 8, MUTED, True)
    set_font(note.add_run("La columna final informa la proporción de períodos en que la fuente publicó una noticia falsa; no es el porcentaje de usuarios ni la proporción total de reposts falsos del sistema."), 8, MUTED)

    add_callout(document, "MAYOR SELECCIÓN NO SIGNIFICA NECESARIAMENTE MÁS FAKE NEWS. ",
                "Las tres combinaciones de mayor selección copian todos los atributos y reducen la tasa de publicaciones falsas a cerca de 10%. La combinación de engagement alcanza menos usuarios (24 por período), pero mantiene 65,9% de publicaciones falsas. Por ello, esta última es la condición más relevante para estudiar la promoción de una fuente que continúa difundiendo fake news.", fill=GOLD)

    document.add_page_break()

    # Página 6: implicaciones, robustez, preguntas y límites.
    add_heading(document, "Implicaciones para una audiencia académica diversa", 1, before=0)
    add_lead_paragraph(document, "Periodismo y estudios de comunicación. ",
                       "Los resultados separan visibilidad, atractivo y veracidad. Una fuente puede adoptar lenguaje simple, recursos audiovisuales, proximidad o sensacionalismo y ganar selecciones sin mejorar la calidad informativa. Para analizar ecosistemas reales conviene medir tanto la circulación como la veracidad de las piezas publicadas.", after=4)
    add_lead_paragraph(document, "Marketing y comportamiento digital. ",
                       "Los atributos de engagement sí modifican la competencia por atención, pero su rendimiento depende del historial del usuario y de las señales sociales. La combinación de 6,03% muestra que optimizar atractivo puede multiplicar el alcance de una fuente falsa, aunque WOM puede contrarrestar ese impulso cuando incorpora una penalización por falsedad.", after=4)
    add_lead_paragraph(document, "Simulación social y ciencias sociales computacionales. ",
                       "Memoria, momento de intervención y red no son parámetros secundarios: forman un mecanismo dinámico. Los efectos no pueden interpretarse aisladamente porque una misma intervención cambia de magnitud según cuánto recuerdan los agentes y cuándo comienza.", after=4)
    add_lead_paragraph(document, "Ingeniería de software y reproducibilidad. ",
                       "La ejecución mediante CLI y copias temporales del Excel permite explorar matrices de parámetros sin modificar el núcleo Java. Las semillas comunes permiten comparar condiciones sobre una base aleatoria compartida y facilitan detectar si el orden de resultados se mantiene.", after=5)

    add_heading(document, "Robustez de los resultados", 2, before=4)
    target_rho_text = f"{target_rho:.3f}".replace(".", ",")
    fake_rho_text = f"{fake_rho:.3f}".replace(".", ",")
    add_lead_paragraph(document, "Los experimentos emparejados con once semillas reprodujeron el orden general. ",
                       f"La correlación de rangos fue rho={target_rho_text} para participación de la fuente y rho={fake_rho_text} para reposts falsos. El emparejamiento redujo la incertidumbre en {narrower} de 10 comparaciones, aunque no lo hizo en dos condiciones con memoria 25.", after=4)
    add_lead_paragraph(document, "Esto fortalece la conclusión comparativa, no la convierte en evidencia empírica. ",
                       "Las semillas muestran que el orden observado no depende principalmente de una realización aleatoria particular. No demuestran que los mismos tamaños de efecto aparezcan en plataformas o poblaciones reales.", after=5)

    add_heading(document, "Preguntas abiertas y próximos experimentos", 2, before=4)
    add_lead_paragraph(document, "Calibración empírica. ", "Usar datos observados de exposición, repost, credibilidad y horizonte temporal para estimar memoria, escala de WOM y base exponencial.", after=2, size=9.0)
    add_lead_paragraph(document, "Mecanismos de red. ", "Incorporar homofilia y comprobar si la penalización producida por WOM persiste cuando los usuarios tienden a conectarse con personas similares.", after=2, size=9.0)
    add_lead_paragraph(document, "Alcance y contacto. ", "Ampliar la matriz de grado, probabilidad de conocer fuentes y alcance para determinar dónde cambia el signo o la magnitud del efecto.", after=2, size=9.0)
    add_lead_paragraph(document, "Validez externa. ", "Reservar parte de los datos empíricos para validación fuera de muestra y evitar que la calibración solo reproduzca el conjunto utilizado para ajustar el modelo.", after=4, size=9.0)

    add_callout(document, "LÍMITES. ", "Los resultados son contrastes descriptivos de una simulación y no estimaciones causales sobre personas reales. Los intervalos son aproximados y no están corregidos por comparaciones múltiples. El modelo aún no incluye homofilia y la calibración empírica permanece pendiente.", fill=GOLD)

    source = document.add_paragraph()
    source.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    source.paragraph_format.space_before = Pt(4)
    source.paragraph_format.space_after = Pt(0)
    set_font(source.add_run(f"Base analítica: {len(runs)} ejecuciones, 274 condiciones, 400 agentes, 400 períodos y 11 repeticiones/semillas por condición. Ventana principal: períodos 301-400."), 7.6, MUTED)

    core = document.core_properties
    core.title = "Síntesis ampliada de experimentos FAKENEWS-ABM"
    core.subject = "Escenarios, memoria, WOM, selección de la fuente y robustez"
    core.author = "FAKENEWS-ABM"
    core.keywords = "fake news, simulación social, WOM, memoria, escenarios"
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    print(f"Wrote {output}")


if __name__ == "__main__":
    main()
